import argparse
import os
import sys
from pathlib import Path

from docx import Document

try:
    import win32com.client as win32  # type: ignore
except Exception:
    win32 = None


def create_filled_word(template_path: str, output_word_path: str, replacements: dict[str, str]) -> None:
    doc = Document(template_path)
    # Replace placeholders in paragraphs
    for p in doc.paragraphs:
        for key, val in replacements.items():
            if key in p.text:
                inline = p.runs
                for i in range(len(inline)):
                    inline[i].text = inline[i].text.replace(key, val)
    # Replace in tables as well
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, val in replacements.items():
                    if key in cell.text:
                        for para in cell.paragraphs:
                            for run in para.runs:
                                run.text = run.text.replace(key, val)
    Path(output_word_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_word_path)


def convert_word_to_pdf(word_path: str, pdf_path: str) -> None:
    if win32 is None:
        raise RuntimeError("pywin32 not available; Word automation requires Windows + pywin32")
    word = win32.Dispatch("Word.Application")
    word.Visible = False
    try:
        doc = word.Documents.Open(os.path.abspath(word_path))
        # 17 = wdFormatPDF
        doc.SaveAs2(os.path.abspath(pdf_path), FileFormat=17)
        doc.Close()
    finally:
        word.Quit()


def main() -> None:
    parser = argparse.ArgumentParser(description="Fill a DOCX template and export as PDF (Windows)")
    parser.add_argument("--template", required=True, help="Path to DOCX template")
    parser.add_argument("--outdir", required=True, help="Output directory for files")
    parser.add_argument("--name", required=True)
    parser.add_argument("--hotel", required=True)
    parser.add_argument("--start", required=True)
    parser.add_argument("--end", required=True)

    args = parser.parse_args()

    name_heb = args.name
    dynamic_pdf_name = f"אישור הזמנה - {name_heb}.pdf"
    dynamic_word_name = f"אישור הזמנה - {name_heb}.docx"

    outdir = os.path.abspath(args.outdir)
    os.makedirs(outdir, exist_ok=True)

    output_word_path = os.path.join(outdir, dynamic_word_name)
    output_pdf_path = os.path.join(outdir, dynamic_pdf_name)

    # Replace placeholders using single-brace keys: {NAME}, {HOTEL}, {START}, {END}
    # (Optionally keep backward-compat for double braces)
    replacements = {
        "{NAME}": args.name,
        "{HOTEL}": args.hotel,
        "{START}": args.start,
        "{END}": args.end,
    }

    create_filled_word(os.path.abspath(args.template), output_word_path, replacements)
    convert_word_to_pdf(output_word_path, output_pdf_path)

    print(f"Saved: {output_pdf_path}")


if __name__ == "__main__":
    try:
        main()
    except Exception as e:
        print(f"ERROR: {e}")
        sys.exit(1)
